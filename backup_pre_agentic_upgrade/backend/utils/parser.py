"""
ScholARA — Document Parser
Extracts raw text from PDF, DOCX, and TXT files
"""

import fitz  # PyMuPDF
import docx
from pathlib import Path
from loguru import logger


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    try:
        doc = fitz.open(file_path)
        text_parts = []

        for page_num, page in enumerate(doc):
            text = page.get_text("text")
            if text.strip():
                # Add page marker for citation purposes
                text_parts.append(f"[Page {page_num + 1}]\n{text}")

        doc.close()
        full_text = "\n\n".join(text_parts)
        logger.info(f"PDF extracted: {len(full_text)} characters from {Path(file_path).name}")
        return full_text

    except Exception as e:
        logger.error(f"PDF extraction error for {file_path}: {e}")
        raise RuntimeError(f"Failed to parse PDF: {e}")


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    try:
        document = docx.Document(file_path)
        paragraphs = []

        for para in document.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract tables
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        full_text = "\n\n".join(paragraphs)
        logger.info(f"DOCX extracted: {len(full_text)} characters from {Path(file_path).name}")
        return full_text

    except Exception as e:
        logger.error(f"DOCX extraction error for {file_path}: {e}")
        raise RuntimeError(f"Failed to parse DOCX: {e}")


def extract_text_from_txt(file_path: str) -> str:
    """Read plain text file."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        logger.info(f"TXT extracted: {len(text)} characters from {Path(file_path).name}")
        return text
    except Exception as e:
        logger.error(f"TXT extraction error for {file_path}: {e}")
        raise RuntimeError(f"Failed to read TXT: {e}")


def parse_document(file_path: str) -> str:
    """
    Main entry point: detect file type and extract text.
    Returns the full extracted text string.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    logger.info(f"Parsing document: {path.name} (type: {suffix})")

    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    elif suffix == ".txt":
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}. Supported: PDF, DOCX, TXT")


def get_document_metadata(file_path: str) -> dict:
    """Extract basic metadata from a PDF (title, author, etc.)."""
    path = Path(file_path)
    meta = {"filename": path.name, "file_type": path.suffix.lower()}

    if path.suffix.lower() == ".pdf":
        try:
            doc = fitz.open(file_path)
            pdf_meta = doc.metadata
            meta.update({
                "title": pdf_meta.get("title", "") or path.stem,
                "author": pdf_meta.get("author", ""),
                "page_count": len(doc),
            })
            doc.close()
        except Exception:
            meta["title"] = path.stem

    return meta
